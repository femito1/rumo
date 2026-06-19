"""Read the guide docx and dump as text/markdown."""
import sys
from docx import Document
from pathlib import Path

src = Path("/home/nandoravioli/bia4u/rumo/Guia_entendimento_workbook_MBC_02_2026.docx")
doc = Document(str(src))

out = []
for p in doc.paragraphs:
    style = p.style.name if p.style else ""
    text = p.text
    if not text.strip():
        out.append("")
        continue
    if style.startswith("Heading"):
        level = style.replace("Heading", "").strip()
        try:
            n = int(level)
        except ValueError:
            n = 2
        out.append("#" * n + " " + text)
    else:
        out.append(text)

# Also dump tables
for ti, table in enumerate(doc.tables):
    out.append(f"\n[TABLE {ti}]")
    for row in table.rows:
        cells = [c.text.replace("\n", " | ") for c in row.cells]
        out.append(" || ".join(cells))

Path("/home/nandoravioli/bia4u/rumo/work/analysis/guide.md").write_text("\n".join(out))
print(f"Wrote {sum(1 for _ in out)} lines, tables={len(doc.tables)}")
